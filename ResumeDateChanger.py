import docx,os
import datetime
# from docx2pdf import convert
doc = docx.Document("------- PATH TO RESUME DOCX FORMAT ONLY ---------")
x=doc.paragraphs[0]
for i in doc.paragraphs:
    print(i.text)
    if( 'Date' in i.text):
        i.clear()
        i.add_run("Date: "+(datetime.date.today().strftime("%d/%m/%Y")+" "+"\n"+"\t\t\t\t\t\t\t\t\t[ <NAME> ]"))
        break
for i in doc.paragraphs:
    print(i.text)
    if(len(i.text.strip())==0 or len(i.text)==0):
        i.clear()
doc.save('------- PATH TO RESUME DOCX FORMAT ONLY ---------')
#Below Line is to using as script
# convert("C:/Users/Pikachu/Downloads/Certs_and_res/Resume.docx","C:/Users/Pikachu/Downloads/Certs_and_res/Resume.pdf")
#Below line is used for making application if used in script format it doesn't execute
os.system("docx2pdf input.docx output.pdf")

